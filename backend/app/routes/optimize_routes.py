import os
import docx
from io import BytesIO
from flask import Blueprint, request, jsonify, send_file, current_app
# We will use your find_and_replace_in_doc function
from app.utils.file_utils import find_and_replace_in_doc, extract_json_from_response
from app.utils.ai_helpers import client
import json # <-- Make sure you have this import for the new route

optimize_bp = Blueprint('optimize', __name__)

# --- NO CHANGES TO THIS SECTION ---
def insert_paragraph_after(paragraph, text, style):
    """A helper function to insert a new paragraph after a given one."""
    new_p = docx.text.paragraph.Paragraph(paragraph._p.getparent()._new_p(), paragraph._parent)
    if text:
        new_p.add_run(text)
    if style:
        new_p.style = style
    paragraph._p.addnext(new_p._p)
    return new_p

@optimize_bp.route("/optimize_resume", methods=["POST"])
def optimize_resume_route():
    """Optimize resume based on user answers to questions"""
    # THIS ENTIRE FUNCTION REMAINS UNCHANGED
    data = request.get_json()
    jd_text = data.get('jd_text')
    answers = data.get('answers')
    file_id = data.get('file_id')
    old_score = data.get('old_score') 
    
    if not all([jd_text, answers, file_id, old_score is not None]):
        return jsonify({"error": "JD, answers, file_id, and old_score are required."}), 400

    target_score = round(old_score * 1.10)

    file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], file_id)
  
    if not os.path.exists(file_path):
        return jsonify({"error": "Original resume file not found."}), 404
        
    if not file_path.lower().endswith('.docx'):
        return jsonify({"error": "Optimization is for .docx only."}), 400
    
    original_doc = docx.Document(file_path)
    original_resume_text = "\n".join([p.text for p in original_doc.paragraphs if p.text.strip()])
    
    formatted_answers = []
    for q, a in answers.items():
        if a.strip():
            formatted_answers.append(f"Question: {q}\nAnswer: {a}\n")
    
    prompt = f"""
You are an expert Career Strategist. Your mission is to surgically enhance a resume to align with a Job Description. Your suggestions must be strong enough to justify a score increase from {old_score}/100 to at least {target_score}/100.
**Perform these three actions:**
1.  **ADD MISSING SKILLS:** Identify 3-6 critical skills from the Job Description that are missing from the resume's SKILLS section.
2.  **ENHANCE PROJECTS WITH NEW BULLETS:** For 1-3 projects in the 'PROFESSIONAL EXPERIENCE' section, identify a missing key achievement that aligns with the JD. Suggest 3-4 high-impact new bullet point to add to that specific project.
3.  **TRANSFORM EXISTING BULLETS:** Rewrite 2-3 existing bullet points to be more impactful with resume and JD point of view.
**CONTEXT:**
**Job Description:**\n{jd_text}
**Candidate Answers:**\n{"".join(formatted_answers)}
**Original Resume:**\n{original_resume_text}
**REQUIRED JSON OUTPUT FORMAT:**
You MUST return ONLY a valid JSON object with EXACTLY three keys: "skills_to_add", "project_enhancements", and "bullet_point_changes".
- `skills_to_add`: An array of strings for missing skills. Example: `["Hadoop", "React", "Scrum"]` so on..
- `project_enhancements`: An array of objects. For each object, provide:
    - `anchor_bullet`: The FULL TEXT of an existing bullet point inside the project where the new bullet should be added. The new bullet will be inserted AFTER privious one.
    - `new_bullet_to_add`: The text of the new bullet point to insert, starting with the bullet character '•'.
- `bullet_point_changes`: An array of objects for find/replace transformations.
**EXAMPLE JSON OUTPUT: 1**
{{
  "skills_to_add": ["Hadoop", "React", "Scrum"],
  "project_enhancements": [
    {{
      "anchor_bullet": "Follows design patterns (like, MVC) and uses logging techniques for better debugging, along with external config file management for flexible settings.",
      "new_bullet_to_add": "• Applied Scrum and Agile methodologies to manage development sprints, prioritize features, and ensure iterative improvements based on feedback."
    }}
  ],
  "bullet_point_changes": [
    {{
      "find": "Developed the backend for an Online Food Ordering System using Java and JDBC, enabling smooth order and cart management.",
      "replace": "Engineered the service-oriented backend for an online ordering platform using Java and JDBC, focusing on scalable and reliable transaction processing."
    }}
  ]
}}
**EXAMPLE JSON OUTPUT:2**
{{
  "skills_to_add": ["SEO/SEM", "Google Analytics", "HubSpot"],
  "project_enhancements": [
    {{
      "anchor_bullet": "Managed social media content calendar across three platforms (Facebook, Instagram, Twitter).",
      "new_bullet_to_add": "• Analyzed campaign performance data using Google Analytics to identify key trends, leading to a 15% increase in audience engagement in Q3."
    }}
  ],
  "bullet_point_changes": [
    {{
      "find": "Responsible for creating weekly email newsletters for the company.",
      "replace": "Authored and distributed weekly email newsletters to a subscriber base of 10,000+, achieving an average open rate of 25% and a 5% click-through rate."
    }}
  ]
}}
Now, generate the JSON output.
"""
    
    try:
        response = client.chat.completions.create(
            # model="gpt-4o",
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            response_format={"type": "json_object"},
            timeout=120.0
        )
        
        response_text = response.choices[0].message.content
        optimization_data = extract_json_from_response(response_text)
        
        skills_to_add = optimization_data.get("skills_to_add", [])
        project_enhancements = optimization_data.get("project_enhancements", [])
        bullet_changes = optimization_data.get("bullet_point_changes", []) 
        
        if not skills_to_add and not project_enhancements and not bullet_changes:
            return jsonify({"error": "The AI could not generate any optimizations."}), 400

        modified_doc = docx.Document(file_path)
        
        if skills_to_add:
            for p in modified_doc.paragraphs:
                if p.text.strip().upper().startswith("SKILLS"):
                    if p.text.strip() and not p.text.strip().endswith(('|', '.', ',')):
                         p.text += ' |'
                    p.text += " " + " | ".join(skills_to_add)
                    break

        if project_enhancements:
            for enhancement in project_enhancements:
                anchor_text = enhancement.get("anchor_bullet")
                new_bullet_text = enhancement.get("new_bullet_to_add")

                if not anchor_text or not new_bullet_text:
                    continue

                for i, p in enumerate(modified_doc.paragraphs):
                    if anchor_text.strip() in p.text.strip():
                        insert_paragraph_after(p, new_bullet_text, p.style)
                        break

        if bullet_changes:
            changes = bullet_changes 
            modified_doc = find_and_replace_in_doc(modified_doc, changes)
        
        doc_io = BytesIO()
        modified_doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name='Optimized_Resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        current_app.logger.error(f"Optimization failed: {str(e)}")
        current_app.logger.error(f"AI response was: {locals().get('response_text', 'Not available')}")
        return jsonify({"error": f"Optimization failed: {str(e)}"}), 500


# # ========================================================================
# # === REPLACE THE OLD /parse_optimized_resume WITH THIS IMPROVED VERSION ===
# # ========================================================================

# @optimize_bp.route("/parse_optimized_resume", methods=["POST"])
# def parse_optimized_resume():
#     """
#     This is the CORRECT and SIMPLE endpoint for the "Design in Builder" button.
#     It receives the FINAL, already-optimized .docx file from the frontend,
#     extracts its text, and asks the AI to parse that text into JSON.
#     """
#     if 'optimized_resume_file' not in request.files:
#         return jsonify({"error": "Optimized resume file is required."}), 400

#     file = request.files['optimized_resume_file']

#     try:
#         from app.utils.file_utils import extract_text_builder
#         final_optimized_text = extract_text_builder(file, file.filename)
#     except Exception as e:
#         return jsonify({"error": f"Failed to extract text from optimized file: {str(e)}"}), 500
    
#     # --- CHANGE 1: A more detailed JSON structure ---
#     target_json_structure = """
#     {
#         "Name": "Full Name", "jobTitle": "Job Title", "email": "email@example.com", "phone": "123-456-7890", "location": "City, Country", "linkedin": "https://linkedin.com/...", "github": "https://github.com/...",
#         "summary": "Professional summary from the text.",
#         "experience": [{"jobTitle": "...", "company": "...", "startDate": "...", "endDate": "...", "description": "..."}],
#         "internship": [{"jobTitle": "...", "company": "...", "startDate": "...", "endDate": "...", "description": "..."}],
#         "education": [{"degree": "...", "school": "...", "startDate": "...", "endDate": "..."}],
#         "skills": ["..."],
#         "languages": ["..."], "interests": ["..."], "achievements": [{"title": "...", "points": ["..."]}],
#         "projects": [{"title": "...", "startDate": "...", "endDate": "...", "tech": "...", "points": ["..."]}],
#         "certifications": [{"name": "...", "issuer": "...", "Date": "..."}],
#         "other_details": "Any text that doesn't fit into the other categories, like 'References available upon request'."
#     }
#     """

#     # --- CHANGE 2: Stronger, more explicit instructions for the AI ---
#     final_prompt = f"""
# You are a meticulous and exhaustive data extraction engine. Your only task is to parse the following resume text and structure it perfectly into the provided JSON format.

# **CRITICAL RULES:**
# 1.  **EXTRACT EVERYTHING:** You MUST capture every single word and detail from the resume text. Do not omit, summarize, or change any information.
# 2.  **NO DATA LOSS:** If you find a piece of text that does not clearly fit into a specific field (like 'experience' or 'projects'), you MUST place it in the `"other_details"` field. Do not discard any data.
# 3.  **COPY VERBATIM:** All text content must be copied exactly as it appears.

# FINAL RESUME TEXT TO PARSE:
# ---
# {final_optimized_text}
# ---

# REQUIRED JSON STRUCTURE:
# {target_json_structure}

# Now, generate the complete and exhaustive JSON object based on the provided text.
# """

#     try:
#         final_response = client.chat.completions.create(
#             # model="gpt-4o",
#             model="gpt-3.5-turbo",
#             messages=[{"role": "user", "content": final_prompt}],
#             temperature=0.0, # Set temperature to 0 for maximum determinism and accuracy
#             response_format={"type": "json_object"}
#         )
#         final_json_data = json.loads(final_response.choices[0].message.content)
#         return jsonify(final_json_data)
        
#     except Exception as e:
#         current_app.logger.error(f"Final JSON parsing failed: {str(e)}")
#         return jsonify({"error": f"Failed to generate resume JSON for builder: {str(e)}"}), 500



# ========================================================================
# === THIS IS THE NEW, CORRECTED ROUTE FOR THE "DESIGN IN BUILDER" BUTTON ===
# ========================================================================

@optimize_bp.route("/parse_final_resume_to_json", methods=["POST"])
def parse_final_resume_to_json():
    """
    Receives the FINAL, already-optimized .docx file from the frontend.
    Its only job is to extract the text and parse it into JSON.
    This guarantees the text is identical to the downloaded file.
    """
    if 'optimized_resume_file' not in request.files:
        return jsonify({"error": "Optimized resume file is required."}), 400

    file = request.files['optimized_resume_file']

    try:
        # Re-import utility here to be safe
        from app.utils.file_utils import extract_text_builder
        # This gets the text from the FINAL .docx file sent from the frontend
        final_optimized_text = extract_text_builder(file, file.filename)
    except Exception as e:
        return jsonify({"error": f"Failed to extract text from optimized file: {str(e)}"}), 500
    
    # The JSON structure your TemplateBuilder expects
    target_json_structure = """
    {
            "Name": "Full Name",
            "jobTitle": "Job Title",
            "email": "email@example.com",
            "phone": "123-456-7890",
            "location": "City, Country",
            "linkedin": "https://linkedin.com/...",
            "github": "https://github.com/...",
            "summary": "Professional summary...",
            "objective": "Carrer Objectives",
            "experience": [
                {{
                    "jobTitle": "Job Title",
                    "company": "Company Name",
                    "startDate": "Month Year",
                    "endDate": "Month Year",
                    "description": ["Point 1", "Point 2"]
                }}
            ],
            "internship": [
                {{
                    "jobTitle": "Job Title",
                    "company": "Company Name",
                    "startDate": "Month Year",
                    "endDate": "Month Year",
                    "description": ["Point 1", "Point 2"]
                }}
            ],
            "education": [
                {{
                    "degree": "Degree Name",
                    "school": "School Name",
                    "level": "type of education level",
                    "startDate": "Year",
                    "endDate": "Year",
                    "cgpa": "X.XX/4.0"  
                }}
            ],
            "skills": ["Skill1", "Skill2"],
            "languages": ["Language1", "Language2"],
            "interests": ["Interest1", "Interest2"],
            "achievements": [
                {{
                    "title": "Achievement Title",
                    "description":" "
                }}
            ],
            "projects": [
                {{
                    "title": "Project Title",
                    "startDate": "Month Year",
                    "endDate": "Month Year",
                    "tech": "Tools/Tech Stack",
                    "description": ["Point 1", "Point 2"]
                }}
            ],
            
             "certifications": [
                {{
                    "name": "Certification name",
                    "issuer": "name of the issuer",
                    "Date": "Month Year",
                }}
            ]
    }
    """

    final_prompt = f"""
You are an expert data extractor. Your task is to parse the following complete resume text and structure it into a clean JSON object.
Do not invent any information. Extract the data exactly as it appears in the text.

FINAL RESUME TEXT:
---
{final_optimized_text}
---

REQUIRED JSON STRUCTURE:
{target_json_structure}

Now, generate the JSON object based on the provided text.
"""

    try:
        final_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": final_prompt}],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        final_json_data = json.loads(final_response.choices[0].message.content)
        return jsonify(final_json_data)
        
    except Exception as e:
        current_app.logger.error(f"Final JSON parsing failed: {str(e)}")
        return jsonify({"error": f"Failed to generate resume JSON for builder: {str(e)}"}), 500
    
    
    
    

@optimize_bp.route("/generate-or-enhance-summary", methods=["POST"])
def generate_or_enhance_summary():
    """
    Generate or enhance a professional summary based on the user's input.
    - If 'existingSummary' is provided, enhances only that text.
    - If 'existingSummary' is empty, generates a new summary using 'experience', 'skills', and 'projects'.
    Expects a JSON payload with 'existingSummary', 'experience', 'skills', and 'projects' fields.
    Returns a JSON object with a 'summary' field containing the generated or enhanced summary.
    """
    data = request.get_json()
    existing_summary = data.get('existingSummary', '').strip()
    experience = data.get('experience', [])
    skills = data.get('skills', [])
    projects = data.get('projects', [])

    # Construct prompt based on whether existing_summary is provided
    if existing_summary:
        # Enhance mode: Only enhance the existing summary
        prompt = f"""
You are an expert Career Strategist tasked with enhancing an existing professional summary for a resume. The enhanced summary should refine the provided text, maintaining its core message and tone, to make it more polished, impactful, and professional. The summary must be 3-5 sentences long and suitable for a resume.

**CONTEXT:**
**Existing Summary:**
{existing_summary}

**REQUIREMENTS:**
- Enhance the existing summary to be more concise, professional, and impactful.
- Maintain the original message and tone, improving clarity and word choice.
- Do not add new information beyond what is provided in the existing summary.
- Keep it concise (3-5 sentences, max 150 words).
- Use a professional tone suitable for a resume.
- Return the result in JSON format with a single key: `summary`.

**EXAMPLE OUTPUT:**
{{
  "summary": "Accomplished Software Engineer with over 5 years of experience, delivering high-quality web applications. Enhanced expertise in modern frameworks, driving system efficiency through scalable solutions."
}}
"""
    else:
        # Generate mode: Create a new summary using experience, skills, and projects
        # Validate input
        if not any([experience, skills, projects]):
            return jsonify({"error": "At least one of experience, skills, or projects is required when generating a new summary."}), 400

        # Format experience for the prompt
        formatted_experience = []
        for exp in experience:
            formatted_experience.append(
                f"Role: {exp.get('jobTitle', '')}\n"
                f"Company: {exp.get('company', '')}\n"
                f"Duration: {exp.get('startDate', '')} - {exp.get('endDate', 'Present')}\n"
                f"Responsibilities: {exp.get('description', '')}\n"
            )

        # Format skills
        formatted_skills = ", ".join(skills) if skills else "None provided"

        # Format projects
        formatted_projects = []
        for proj in projects:
            formatted_projects.append(
                f"Project: {proj.get('title', '')}\n"
                f"Duration: {proj.get('startDate', '')} - {proj.get('endDate', 'Present')}\n"
                f"Description: {proj.get('description', '')}\n"
                f"Technologies: {proj.get('tech', '')}\n"
            )

        prompt = f"""
You are an expert Career Strategist tasked with generating a concise, impactful professional summary for a resume. The summary should highlight the candidate's experience, skills, and projects, tailored to a professional context. The summary must be 3-5 sentences long, professional, and aligned with the provided context.

**CONTEXT:**
**Experience:**
{"".join(formatted_experience) if formatted_experience else "None provided"}

**Skills:**
{formatted_skills}

**Projects:**
{"".join(formatted_projects) if formatted_projects else "None provided"}

**REQUIREMENTS:**
- Generate a professional summary that reflects the candidate's experience, skills, and projects.
- Keep it concise (3-5 sentences, max 150 words).
- Use a professional tone suitable for a resume.
- Incorporate key achievements or responsibilities from experience and projects.
- Highlight relevant skills to showcase expertise.
- Return the result in JSON format with a single key: `summary`.

**EXAMPLE OUTPUT:**
{{
  "summary": "Results-driven Software Engineer with over 5 years of experience developing scalable web applications at leading tech firms. Proficient in Java, Python, and cloud technologies, with a proven track record of delivering high-impact projects that enhance system performance. Skilled in leading cross-functional teams to meet tight deadlines while maintaining code quality."
}}
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,  # Moderate temperature for balanced creativity and consistency
            response_format={"type": "json_object"},
            timeout=60.0
        )
        
        response_text = response.choices[0].message.content
        summary_data = json.loads(response_text)
        
        if not summary_data.get("summary"):
            return jsonify({"error": "Failed to generate a valid summary."}), 500

        return jsonify(summary_data)
        
    except OpenAIError as e:
        current_app.logger.error(f"OpenAI API error during summary generation/enhancement: {str(e)}")
        return jsonify({"error": f"Failed to process summary: {str(e)}"}), 500
    except json.JSONDecodeError as e:
        current_app.logger.error(f"JSON parsing error: {str(e)}")
        return jsonify({"error": "Invalid response format from AI model."}), 500
    except Exception as e:
        current_app.logger.error(f"Unexpected error during summary generation/enhancement: {str(e)}")
        return jsonify({"error": f"Failed to process summary: {str(e)}"}), 500
